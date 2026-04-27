import os
import sys
import json
import smtplib
import subprocess
import urllib.request
import urllib.error
import time
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document as ReadDoc
from docx import Document as WriteDoc
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))

# ==================== CONFIG ====================
USER_NAME        = os.getenv("USER_NAME", "Amine Ouardi")
USER_EMAIL       = os.getenv("USER_EMAIL", "aminerc.business@gmail.com")
OUTPUT_FOLDER    = os.path.normpath(os.getenv("OUTPUT_FOLDER", "."))
EMAIL_FROM       = os.getenv("EMAIL_FROM", "")
EMAIL_TO         = os.getenv("EMAIL_TO", "")
EMAIL_PASSWORD   = os.getenv("EMAIL_PASSWORD", "")
SMTP_HOST        = os.getenv("SMTP_HOST", "")
SMTP_PORT        = int(os.getenv("SMTP_PORT", "0") or "0")
SMTP_MODE        = os.getenv("SMTP_MODE", "").lower().strip()  # "ssl" | "starttls" | ""
SMTP_USER        = os.getenv("SMTP_USER", "")

USE_GEMINI       = os.getenv("USE_GEMINI", "true").lower() == "true"
USE_CLAUDE       = os.getenv("USE_CLAUDE", "false").lower() == "true"

GEMINI_API_KEY   = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL     = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")
CLAUDE_API_KEY   = os.getenv("CLAUDE_API_KEY", "")
CLAUDE_MODEL     = os.getenv("CLAUDE_MODEL", "claude-haiku-4-5-20251001")

MAX_RETRIES      = 5

# ==================== NOTIFICATION WINDOWS ====================
def _ensure_windows_toast_app_id(app_id: str):
    """
    Windows peut ignorer les toasts si l'AppUserModelID n'est pas "enregistré" via un raccourci Start Menu.
    On crée (si absent) un .lnk avec l'AppID pour rendre les notifications fiables.
    """
    try:
        start_menu = os.path.join(
            os.environ.get("APPDATA", ""),
            "Microsoft",
            "Windows",
            "Start Menu",
            "Programs",
        )
        if not start_menu or not os.path.isdir(start_menu):
            return

        shortcut_path = os.path.join(start_menu, "meeting-sumup.lnk")
        if os.path.isfile(shortcut_path):
            return

        py_exe = sys.executable
        # Point d'entrée "inoffensif" : on associe juste l'AppID à une commande python.
        # (Le toast utiliserait ensuite CreateToastNotifier(app_id).)
        args = '-c "import sys; sys.exit(0)"'

        ps = rf"""
$ErrorActionPreference = "Stop"
$shortcutPath = "{shortcut_path}"
$targetPath   = "{py_exe}"
$arguments    = '{args}'
$workingDir   = "{os.path.dirname(os.path.abspath(__file__))}"
$appId        = "{app_id}"

$WshShell = New-Object -ComObject WScript.Shell
$Shortcut = $WshShell.CreateShortcut($shortcutPath)
$Shortcut.TargetPath = $targetPath
$Shortcut.Arguments  = $arguments
$Shortcut.WorkingDirectory = $workingDir
$Shortcut.WindowStyle = 7
$Shortcut.Save()

$code = @"
using System;
using System.Runtime.InteropServices;

[ComImport, Guid("886D8EEB-8CF2-4446-8D02-CDBA1DBDCF99"), InterfaceType(ComInterfaceType.InterfaceIsIUnknown)]
interface IPropertyStore {{
    uint GetCount(out uint cProps);
    uint GetAt(uint iProp, out PROPERTYKEY pkey);
    uint GetValue(ref PROPERTYKEY key, out PROPVARIANT pv);
    uint SetValue(ref PROPERTYKEY key, ref PROPVARIANT pv);
    uint Commit();
}}

[StructLayout(LayoutKind.Sequential, Pack = 4)]
struct PROPERTYKEY {{
    public Guid fmtid;
    public uint pid;
}}

[StructLayout(LayoutKind.Sequential)]
struct PROPVARIANT {{
    public ushort vt;
    public ushort wReserved1;
    public ushort wReserved2;
    public ushort wReserved3;
    public IntPtr p;
    public int p2;
}}

static class NativeMethods {{
    [DllImport("shell32.dll", CharSet = CharSet.Unicode)]
    public static extern int SHGetPropertyStoreFromParsingName(
        [MarshalAs(UnmanagedType.LPWStr)] string pszPath,
        IntPtr zeroWorks,
        uint flags,
        ref Guid riid,
        out IPropertyStore propertyStore);

    [DllImport("ole32.dll")]
    public static extern int PropVariantClear(ref PROPVARIANT pvar);
}}

public class ShortcutAumid {{
    static PROPERTYKEY PKEY_AppUserModel_ID = new PROPERTYKEY {{
        fmtid = new Guid("9F4C2855-9F79-4B39-A8D0-E1D42DE1D5F3"),
        pid = 5
    }};

    public static void SetAppId(string shortcut, string appId) {{
        Guid guid = new Guid("886D8EEB-8CF2-4446-8D02-CDBA1DBDCF99");
        IPropertyStore store;
        int hr = NativeMethods.SHGetPropertyStoreFromParsingName(shortcut, IntPtr.Zero, 0, ref guid, out store);
        if (hr != 0) Marshal.ThrowExceptionForHR(hr);

        PROPVARIANT pv = new PROPVARIANT();
        pv.vt = 31; // VT_LPWSTR
        pv.p = Marshal.StringToCoTaskMemUni(appId);
        try {{
            store.SetValue(ref PKEY_AppUserModel_ID, ref pv);
            store.Commit();
        }} finally {{
            NativeMethods.PropVariantClear(ref pv);
        }}
    }}
}}
"@

Add-Type -TypeDefinition $code -Language CSharp
[ShortcutAumid]::SetAppId($shortcutPath, $appId)
"""
        # Lancer PowerShell en mode silencieux; si ça échoue, on n'empêche pas le reste.
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        subprocess.run(
            ["powershell.exe", "-ExecutionPolicy", "Bypass", "-NoProfile", "-Command", ps],
            check=False,
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            startupinfo=si,
        )
    except Exception:
        return

def notify_windows(meeting_name, date_str, summary, output_path):
    """
    Notification Windows (fallback console si plyer indisponible).
    """
    force_popup = os.getenv("NOTIFY_FORCE_POPUP", "false").lower().strip() == "true"
    # plyer (balloontip) passe par NOTIFYICONDATAW avec des limites strictes.
    # Si on depasse, l'exception est levee dans un thread interne (difficile a catcher ici),
    # donc on tronque agressivement avant l'appel.
    title = "Synthese generée"[:63]
    abs_path = os.path.abspath(output_path) if output_path else ""
    msg = "La synthese de la réunion a été generée".replace("\r", " ").strip()[:240]
    app_name = "meeting-sumup"[:63]

    print(f"[NOTIF] {output_path}")

    # 1) Toast natif via PowerShell/WinRT (sans image, comme ton test qui marche)
    try:
        _ensure_windows_toast_app_id(app_name)
        file_url = ""
        if abs_path:
            file_url = "file:///" + abs_path.replace("\\", "/")

        ps = rf'''
[Windows.UI.Notifications.ToastNotificationManager, Windows.UI.Notifications, ContentType = WindowsRuntime] | Out-Null
[Windows.UI.Notifications.ToastNotification, Windows.UI.Notifications, ContentType = WindowsRuntime] | Out-Null
[Windows.Data.Xml.Dom.XmlDocument, Windows.Data.Xml.Dom.XmlDocument, ContentType = WindowsRuntime] | Out-Null

$xml = New-Object Windows.Data.Xml.Dom.XmlDocument
$xml.LoadXml(@"
<toast activationType="protocol" launch="{file_url}">
  <visual>
    <binding template="ToastText02">
      <text id="1"><![CDATA[{title}]]></text>
      <text id="2"><![CDATA[{msg}]]></text>
    </binding>
  </visual>
</toast>
"@)

$toast = [Windows.UI.Notifications.ToastNotification]::new($xml)
$n = [Windows.UI.Notifications.ToastNotificationManager]::CreateToastNotifier("{app_name}")
$n.Show($toast)
'''

        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        subprocess.Popen(
            ["powershell.exe", "-ExecutionPolicy", "Bypass", "-NoProfile", "-Command", ps],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            startupinfo=si,
        )
        if force_popup:
            raise RuntimeError("NOTIFY_FORCE_POPUP=true")
        return
    except Exception as e:
        print(f"[WARN] Toast PowerShell indisponible ({e}). Fallback plyer...")

    try:
        from plyer import notification  # type: ignore
        notification.notify(
            title=title,
            message=msg,
            timeout=12,
            app_name=app_name,
        )
        if force_popup:
            raise RuntimeError("NOTIFY_FORCE_POPUP=true")
    except Exception as e:
        # Fallback visible (bloquant) si tout le reste est bloque/desactive
        try:
            import ctypes

            ctypes.windll.user32.MessageBoxW(0, msg, title, 0x00001000)  # MB_SYSTEMMODAL
            return
        except Exception:
            pass

        print(f"[WARN] Notification Windows indisponible ({e}).")
        print(title)
        print(msg)

# ==================== LECTURE TRANSCRIPTION ====================
def read_transcript(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        doc = ReadDoc(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext in (".txt", ".vtt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

# ==================== PROMPT ====================
def build_prompt(transcript):
    return f"""Tu es un assistant expert en prise de notes de reunions professionnelles.
L'utilisateur principal s'appelle {USER_NAME}.

Voici la transcription d'un meeting :

---
{transcript[:30000]}
---

Genere une synthese structuree en francais avec exactement ce format :

## TITRE
[5 mots max resumant le sujet principal du meeting]

## RESUME
[2-4 phrases resumant le contexte et les sujets abordes]

## DECISIONS PRISES
- [Decision 1]
[Si aucune : "Aucune decision formelle identifiee"]

## ACTIONS DE {USER_NAME.upper()}
- [Action concernant {USER_NAME} specifiquement] - [Deadline si mentionnee]
[Si aucune : "Aucune action identifiee pour {USER_NAME}"]

## ACTIONS GENERALES
- [Action] - [Responsable] - [Deadline si mentionnee]
[Si aucune : "Aucune action formelle identifiee"]

## POINTS CLES
- [Point important 1]
- [Point important 2]
"""

# ==================== APPEL API ====================
def call_gemini(prompt):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = json.dumps({"contents": [{"parts": [{"text": prompt}]}]}).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"}, method="POST")

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            with urllib.request.urlopen(req) as response:
                data = json.loads(response.read().decode())
                return data["candidates"][0]["content"]["parts"][0]["text"]
        except urllib.error.HTTPError as e:
            if e.code != 429:
                raise
            wait_s = min(60, 2 ** (attempt - 1))
            print(f"Rate limit. Attente {wait_s}s...")
            time.sleep(wait_s)
    raise RuntimeError("Gemini rate limit depasse.")

def call_claude(prompt):
    url = "https://api.anthropic.com/v1/messages"
    payload = json.dumps({
        "model": CLAUDE_MODEL,
        "max_tokens": 2048,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={
        "Content-Type": "application/json",
        "x-api-key": CLAUDE_API_KEY,
        "anthropic-version": "2023-06-01"
    }, method="POST")
    with urllib.request.urlopen(req) as response:
        data = json.loads(response.read().decode())
        return data["content"][0]["text"]

def summarize(transcript):
    prompt = build_prompt(transcript)
    if USE_CLAUDE and CLAUDE_API_KEY:
        print("Appel Claude...")
        return call_claude(prompt)
    elif USE_GEMINI and GEMINI_API_KEY:
        print("Appel Gemini...")
        return call_gemini(prompt)
    else:
        raise RuntimeError("Aucune API active ou configuree.")

# ==================== CREATION DOCX ====================
def create_docx(summary, meeting_name, date_str, output_path):
    from docx import Document
    from docx.shared import Pt
    import tempfile

    doc = Document()

    h = doc.add_heading(f"Notes de reunion — {meeting_name}", level=1)
    h.runs[0].font.size = Pt(16)

    doc.add_paragraph(f"Date : {date_str}")
    doc.add_paragraph(f"Fichier source : {meeting_name}")
    doc.add_paragraph("")

    for line in summary.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        else:
            doc.add_paragraph(line)

    # Robustesse: s'assurer que le dossier de sortie existe (OneDrive / run context)
    parent_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(parent_dir, exist_ok=True)

    try:
        doc.save(output_path)
        print(f"Fichier Word cree : {output_path}")
        return output_path
    except FileNotFoundError as e:
        # Cas observe avec certains dossiers OneDrive "placeholder" (reparse point) non ecrivable
        fallback_dir = os.path.join(tempfile.gettempdir(), "meeting-notes")
        os.makedirs(fallback_dir, exist_ok=True)
        fallback_path = os.path.join(fallback_dir, os.path.basename(output_path))
        doc.save(fallback_path)
        print(f"[ERREUR] Impossible d'ecrire dans : {output_path}")
        print(f"[INFO] Fichier Word cree en fallback : {fallback_path}")
        raise RuntimeError(
            f"{e}\n"
            f"Dossier de sortie non ecrivable (souvent OneDrive placeholder).\n"
            f"Fix: rendre le dossier '{parent_dir}' disponible hors-ligne / recreer le dossier,\n"
            f"ou changer OUTPUT_FOLDER dans .env.\n"
            f"Fallback cree ici: {fallback_path}"
        )

# ==================== EMAIL ====================
def send_email(meeting_name, date_str, summary, output_path):
    email_from = (EMAIL_FROM or USER_EMAIL).strip()
    email_to_raw = (EMAIL_TO or USER_EMAIL).strip()
    email_password = (EMAIL_PASSWORD or "").strip()

    if not email_from or not email_to_raw:
        raise RuntimeError(
            "Email non configure: definir EMAIL_FROM et EMAIL_TO (ou USER_EMAIL) dans .env."
        )
    if not email_password:
        raise RuntimeError(
            "Email non configure: definir EMAIL_PASSWORD dans .env (mot de passe/app password SMTP)."
        )

    # Support: EMAIL_TO peut contenir plusieurs emails separes par virgules/points-virgules
    recipients = [x.strip() for x in email_to_raw.replace(";", ",").split(",") if x.strip()]
    if not recipients:
        raise RuntimeError("EMAIL_TO invalide: aucun destinataire detecte.")

    subject = f"Notes meeting — {meeting_name} — {date_str}"
    body = f"""Bonjour {USER_NAME},

Les notes de ta reunion sont disponibles ici :
{output_path}

---
APERCU :

{summary[:1500]}
---

Notes completes dans le fichier Word ci-dessus.
"""
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"]    = email_from
    msg["To"]      = ", ".join(recipients)
    msg.attach(MIMEText(body, "plain"))

    # Attacher le docx si disponible
    try:
        if output_path and os.path.isfile(output_path):
            with open(output_path, "rb") as f:
                part = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(output_path)}"')
            msg.attach(part)
    except Exception as e:
        print(f"[WARN] Impossible d'attacher le fichier: {e}")

    def _infer_smtp(from_addr: str):
        domain = (from_addr.split("@")[-1] if "@" in from_addr else "").lower()
        if domain in {"outlook.com", "hotmail.com", "live.com", "msn.com"}:
            return ("smtp.office365.com", 587, "starttls")
        if domain in {"icloud.com", "me.com", "mac.com"}:
            return ("smtp.mail.me.com", 587, "starttls")
        if domain in {"gmail.com", "googlemail.com"}:
            return ("smtp.gmail.com", 465, "ssl")
        return ("smtp.gmail.com", 587, "starttls")

    host = SMTP_HOST.strip() or _infer_smtp(email_from)[0]
    port = SMTP_PORT or _infer_smtp(email_from)[1]
    mode = SMTP_MODE or _infer_smtp(email_from)[2]
    user = (SMTP_USER.strip() or email_from)

    if mode == "ssl":
        with smtplib.SMTP_SSL(host, port) as server:
            server.login(user, email_password)
            server.sendmail(email_from, recipients, msg.as_string())
    else:
        with smtplib.SMTP(host, port) as server:
            server.ehlo()
            if mode == "starttls":
                server.starttls()
                server.ehlo()
            server.login(user, email_password)
            server.sendmail(email_from, recipients, msg.as_string())

    print(f"Email envoye a {', '.join(recipients)}")

# ==================== MAIN ====================
def process(file_path):
    print(f"\n[{datetime.now().strftime('%H:%M:%S')}] Traitement : {file_path}")

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    file_name   = os.path.basename(file_path)
    date_str    = datetime.now().strftime("%Y-%m-%d %H:%M")
    date_slug   = datetime.now().strftime("%Y%m%d_%H%M")

    # 1. Lecture
    transcript = read_transcript(file_path)
    print(f"Transcription lue : {len(transcript)} caracteres")

    # 2. Synthese
    summary = summarize(transcript)

    # 3. Extraire le titre
    title = "meeting"
    prev_line = ""
    for line in summary.split("\n"):
        if line.strip() and not line.startswith("##"):
            prev_line = ""
        if "## TITRE" in line:
            prev_line = line
        elif prev_line and "## TITRE" in prev_line:
            title = line.strip().replace(" ", "_").replace("/", "-")
            break

    # Nettoyer le titre pour nom de fichier
    import re
    title = re.sub(r"[^\w\-]", "_", title)[:50]
    output_name = f"notes_{title}_{date_slug}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, output_name)

    # 4. Word
    create_docx(summary, file_name, date_str, output_path)

    # 5. Notification Windows (remplace l'email)
    notify_windows(file_name, date_str, summary, output_path)

    print(f"[OK] {output_name}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage : python process.py <fichier_transcription>")
        sys.exit(1)
    process(sys.argv[1])
